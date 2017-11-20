# -*- coding: utf-8 -*-
'''
按月份整合word至excel
'''

from docx import Document
from pandas import DataFrame
from openpyxl import Workbook
import os

def getpathlist():
    y = ''
    while not y == 'y':
        y = input('请确保该脚本所处目录下仅包括当月docx文件，其他格式文件没有关系\n继续请输入y')
    rootdir = os.getcwd()
    path_list = []
    
    
    for parent, dirnames,filenames in os.walk(rootdir):
        for fn in filenames:
            if fn.endswith('.docx') and date in fn and '住院运行病历质量考核表' in fn :
                print('fn:{}'.format(fn))
                path_list.append(os.path.join(parent,fn))
    return path_list
    
def getdoclist(path_list):
    doc_list=[]
    for i in path_list:
        with open(i,'rb') as f:
            document = Document(f)
            doc_list.append(document)
    return doc_list
'''
    for p in path_list:
        try:
            f = open(path_list,'rb')
            document = Document(f)
            doc_list.append(document)
        
        except Exception as e:
            print('exception happend:{}'.format(e))
        
        finally:
            f.close()
'''
def getinfo(doc_list):
    wb =  Workbook()
    ws = wb.active
    ws.merge_cells('A1:F1')
    
    ws['A1'].value = date
    ws['A2'].value ='科室'
    ws['B2'].value = '病房'
    ws['C2'].value = '患者姓名'
    ws['D2'].value = '病历号'
    ws['E2'].value = '入院时间'
    ws['F2'].value = '出院时间'
    
    nrow = 3
    
    for d in doc_list: #遍历每个文件
        
        ks_info = d.paragraphs[1].text #获取抬头
        print('ks_info:{}'.format(ks_info))
        try:
            ks = ks_info.split()[1] #科室名称
        except:
            ks = None
        try:
            bf = ks_info.split()[3] #病房名称
        except:
            bf = None
        table = d.tables[0]
        

        for col in range(-5,0): #遍历列

            hzxm = table.column_cells(col)[-7].text #患者姓名
            zyh = table.column_cells(col)[-6].text #住院号
            rysj = table.column_cells(col)[-5].text #入院时间
            ws['A{}'.format(nrow)].value = ks
            ws['B{}'.format(nrow)].value = bf
            ws['C{}'.format(nrow)].value = hzxm
            ws['D{}'.format(nrow)].value = zyh
            ws['E{}'.format(nrow)].value = rysj
            nrow += 1
    
    
    wb.save('{}.xlsx'.format(date))

def main():
    
    pathlist = getpathlist()
    doclist = getdoclist(pathlist)
    getinfo(doclist)
    
date = input('请输入日期，按照文件命名格式')    
main()